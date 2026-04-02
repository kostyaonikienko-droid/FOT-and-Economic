import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .db_manager import get_setting, get_spirit_norm
from .grouping import GROUPS, determine_group

THIN_BORDER = Border(left=Side(style='thin'),
                     right=Side(style='thin'),
                     top=Side(style='thin'),
                     bottom=Side(style='thin'))

def format_cell(cell, font_size=10, bold=False, alignment='center', wrap_text=False):
    cell.font = Font(size=font_size, bold=bold)
    cell.alignment = Alignment(horizontal=alignment, vertical='center', wrap_text=wrap_text)
    cell.border = THIN_BORDER

def get_employee_name(fot_widget, tab_number: str) -> str:
    """Получить ФИО по табельному номеру через fot_widget"""
    if not fot_widget or not tab_number:
        return ''
    if hasattr(fot_widget, 'project') and fot_widget.project:
        for emp in fot_widget.project.employees:
            if str(emp.tab_num) == str(tab_number):
                return emp.fio
    return ''

def generate_act_word(works_by_group: Dict[str, List[Dict]], period_start: str, period_end: str, output_path: str, fot_widget=None):
    """Генерация акта в формате Word (docx) по образцу Пример.docx"""
    doc = Document()
    
    # 1. Строка с антисептиком и инв. №
    antiseptic_info = get_setting("antiseptic_solution_ticket") or ""
    inv_number = get_setting("antiseptic_inv_number") or ""
    p = doc.add_paragraph()
    p.add_run(f"Раствор антисептический получен по тр. {antiseptic_info}").bold = True
    p.add_run(f"\tИнв. № {inv_number}")
    
    # 2. Период
    doc.add_paragraph(f"с {period_start} по {period_end}")
    doc.add_paragraph()  # пустая строка
    
    # 3. Группы с таблицами
    for group_full_name, works in works_by_group.items():
        if not works:
            continue
        # Заголовок группы (жирный, как в примере)
        heading = doc.add_paragraph(group_full_name)
        heading.runs[0].bold = True
        
        # Таблица 3x3 (9 колонок)
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i in range(3):
            hdr_cells[i*3].text = "Дата"
            hdr_cells[i*3+1].text = "№ счета"
            hdr_cells[i*3+2].text = "к-во"
        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Распределение работ по трём колонкам
        columns_data = [[], [], []]
        for i, w in enumerate(works):
            col_idx = i % 3
            columns_data[col_idx].append((w['date_work'], w['invoice_num'], w.get('quantity', 1)))
        max_rows = max(len(col) for col in columns_data) if columns_data else 0
        for row_idx in range(max_rows):
            row_cells = table.add_row().cells
            for col_idx in range(3):
                if row_idx < len(columns_data[col_idx]):
                    date_val, inv_val, qty_val = columns_data[col_idx][row_idx]
                    row_cells[col_idx*3].text = date_val
                    row_cells[col_idx*3+1].text = inv_val
                    row_cells[col_idx*3+2].text = str(qty_val)
                else:
                    row_cells[col_idx*3].text = ""
                    row_cells[col_idx*3+1].text = ""
                    row_cells[col_idx*3+2].text = ""
        
        # Итог по группе
        total_quantity = sum(w.get('quantity', 1) for w in works)
        group_key = determine_group(works[0]['work_name']) if works else None
        norm = get_spirit_norm(group_key) if group_key else 0
        spisano = total_quantity * norm
        doc.add_paragraph(f"Итого: {total_quantity} шт. расход спирта - {norm:.3f} л/шт Списано (литров) — {spisano:.2f}")
        doc.add_paragraph()  # пустая строка между группами
    
    # 4. Общий итог
    total_spirit = 0
    for group_full_name, works in works_by_group.items():
        if not works:
            continue
        group_key = determine_group(works[0]['work_name']) if works else None
        if group_key:
            norm = get_spirit_norm(group_key)
            total_quantity = sum(w.get('quantity', 1) for w in works)
            total_spirit += total_quantity * norm
    doc.add_paragraph(f"Всего списано спирта — {total_spirit:.2f} литров")
    
    doc.save(output_path)

def generate_act_excel(works_by_group: Dict[str, List[Dict]], period_start: str, period_end: str, output_path: str, fot_widget=None):
    """Генерация акта в Excel (по образцу Пример.xlsx) с расширенным журналом"""
    wb = openpyxl.Workbook()
    ws_act = wb.active
    ws_act.title = "Акт"

    # ---- Шапка (как в примере: нет объединения) ----
    antiseptic_info = get_setting("antiseptic_solution_ticket") or ""
    inv_number = get_setting("antiseptic_inv_number") or ""
    ws_act['A1'] = f"Раствор антисептический получен по тр. {antiseptic_info}"
    format_cell(ws_act['A1'], alignment='left')
    ws_act['H1'] = "Инв. №"
    ws_act['I1'] = inv_number
    format_cell(ws_act['H1'], bold=True)
    format_cell(ws_act['I1'])

    ws_act['A2'] = f"с {period_start} по {period_end}"
    format_cell(ws_act['A2'], alignment='left')

    current_row = 4  # после двух строк и пустой строки (строка 3 пустая)

    # ---- Группы ----
    for group_full_name, works in works_by_group.items():
        if not works:
            continue
        # Заголовок группы
        ws_act.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=9)
        cell = ws_act.cell(row=current_row, column=1)
        cell.value = group_full_name
        format_cell(cell, font_size=10, bold=True, alignment='left')
        current_row += 1

        # Шапка таблицы (три блока)
        headers_row = current_row
        for col_group in range(3):
            start_col = col_group * 3 + 1
            ws_act.cell(row=headers_row, column=start_col, value="Дата")
            ws_act.cell(row=headers_row, column=start_col+1, value="№ счета")
            ws_act.cell(row=headers_row, column=start_col+2, value="к-во")
            for c in range(start_col, start_col+3):
                format_cell(ws_act.cell(row=headers_row, column=c), bold=True)
        current_row += 1

        # Распределение данных по трём колонкам
        columns_data = [[], [], []]
        for i, w in enumerate(works):
            col_idx = i % 3
            columns_data[col_idx].append((w['date_work'], w['invoice_num'], w.get('quantity', 1)))
        max_rows = max(len(col) for col in columns_data)
        for row_offset in range(max_rows):
            for col_idx in range(3):
                if row_offset < len(columns_data[col_idx]):
                    date_val, inv_val, qty_val = columns_data[col_idx][row_offset]
                    base_col = col_idx * 3 + 1
                    cell = ws_act.cell(row=current_row + row_offset, column=base_col)
                    cell.value = date_val
                    format_cell(cell, alignment='center')
                    cell = ws_act.cell(row=current_row + row_offset, column=base_col+1)
                    cell.value = inv_val
                    format_cell(cell, alignment='center')
                    cell = ws_act.cell(row=current_row + row_offset, column=base_col+2)
                    cell.value = qty_val
                    format_cell(cell, alignment='center')
        current_row += max_rows + 1

        # Итог по группе
        total_quantity = sum(w.get('quantity', 1) for w in works)
        group_key = None
        for key, full in GROUPS.items():
            if full == group_full_name:
                group_key = key
                break
        norm = get_spirit_norm(group_key) if group_key else 0
        spisano = total_quantity * norm
        total_line = f"Итого: {total_quantity} шт. расход спирта - {norm:.3f} л/шт Списано (литров) — {spisano:.2f}"
        ws_act.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=9)
        cell = ws_act.cell(row=current_row, column=1)
        cell.value = total_line
        format_cell(cell, bold=True, alignment='left')
        current_row += 2

    # ---- Общий итог ----
    total_spirit = 0
    for group_full_name, works in works_by_group.items():
        if not works:
            continue
        group_key = None
        for key, full in GROUPS.items():
            if full == group_full_name:
                group_key = key
                break
        if group_key:
            norm = get_spirit_norm(group_key)
            total_quantity = sum(w.get('quantity', 1) for w in works)
            total_spirit += total_quantity * norm
    ws_act.merge_cells(start_row=current_row, start_column=1, end_row=current_row, end_column=4)
    cell = ws_act.cell(row=current_row, column=1)
    cell.value = f"Всего списано спирта — {total_spirit:.2f} литров"
    format_cell(cell, bold=True, alignment='left')

    # ---- Журнал (расширенный) ----
    ws_journal = wb.create_sheet("Журнал")
    headers = ["Дата вып", "Дата сч", "№ сч", "Сумма", "Работа", "Кол-во", "Таб. №", "ФИО", "Общее кол-во литров по счету"]
    for col_idx, header in enumerate(headers, 1):
        cell = ws_journal.cell(row=1, column=col_idx)
        cell.value = header
        format_cell(cell, bold=True)

    all_works = []
    for works_list in works_by_group.values():
        all_works.extend(works_list)

    # Вычисляем общее кол-во литров по каждому счету
    invoice_total_liters = {}
    for w in all_works:
        inv = w['invoice_num']
        group_key = determine_group(w['work_name'])
        norm = get_spirit_norm(group_key) if group_key else 0
        liters = w.get('quantity', 1) * norm
        invoice_total_liters[inv] = invoice_total_liters.get(inv, 0) + liters

    row_idx = 2
    for w in all_works:
        ws_journal.cell(row=row_idx, column=1, value=w['date_work'])
        ws_journal.cell(row=row_idx, column=2, value=w['date_invoice'])
        ws_journal.cell(row=row_idx, column=3, value=w['invoice_num'])
        ws_journal.cell(row=row_idx, column=4, value=w['amount'])
        ws_journal.cell(row=row_idx, column=5, value=w['work_name'])
        ws_journal.cell(row=row_idx, column=6, value=w.get('quantity', 1))
        ws_journal.cell(row=row_idx, column=7, value=w.get('tab_number', ''))
        # ФИО
        fio = get_employee_name(fot_widget, w.get('tab_number', ''))
        ws_journal.cell(row=row_idx, column=8, value=fio)
        # Общее кол-во литров по счету
        total_liters = invoice_total_liters.get(w['invoice_num'], 0)
        ws_journal.cell(row=row_idx, column=9, value=total_liters)
        row_idx += 1

    for col in ws_journal.columns:
        max_length = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws_journal.column_dimensions[col_letter].width = adjusted_width

    wb.save(output_path)

# Основная функция
def generate_act(works_by_group: Dict[str, List[Dict]], period_start: str, period_end: str, output_path: str, fot_widget=None, format='word'):
    if format.lower() == 'word':
        generate_act_word(works_by_group, period_start, period_end, output_path, fot_widget)
    else:
        generate_act_excel(works_by_group, period_start, period_end, output_path, fot_widget)